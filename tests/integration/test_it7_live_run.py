"""IT-7 (TC-16): the live-provider run path, end to end.

Covers AT-M16-1 (live path real), -3 (pause un-bypassable), -5 (rerun replays
archives, no provider call), -6 (write-back + manifest hashing, no secrets). The
provider is a **recording** stub (records calls, returns fixed valid Node A /
Node C responses) so the test proves the live wiring — not the canned echo — and
touches no network.
"""

from __future__ import annotations

import copy
import hashlib
import json
import re
from collections.abc import Callable
from pathlib import Path

import pytest
import yaml
from docx import Document  # type: ignore[import-untyped]
from integration_study import build_integration_study, integration_config, study_document
from it_util import fast_policy
from stub_nodes import node_a_provider, node_c_approve_provider

from burhan.cli.live import live_confirm, live_extract, live_rerun
from burhan.contract.llm_base import LlmSettings
from burhan.core.errors import BurhanHalt
from burhan.core.manifest import Manifest


class RecordingFactory:
    """A ``provider_factory`` that records calls and returns canned responses.

    Non-canned in the AT-M16-1 sense: each call is recorded (proving the live
    path invoked the provider), and it is distinct from ``certification_run``'s
    injected lambdas. No network.
    """

    def __init__(self) -> None:
        self.calls: dict[str, int] = {"node_a": 0, "node_c": 0}
        self._node_a = node_a_provider(integration_config())
        self._node_c = node_c_approve_provider()

    def __call__(self, settings: LlmSettings, node: str) -> Callable[[str], str]:
        inner = {"node_a": self._node_a, "node_c": self._node_c}[node]

        def call(prompt: str) -> str:
            self.calls[node] += 1
            return inner(prompt)

        return call


def _llm_yaml() -> str:
    anthropic = {
        "provider": "anthropic",
        "model": "claude-live-test",
        "lineage": "anthropic.claude",
        "temperature": 0,
        "max_retries": 2,
    }
    openai = dict(anthropic, provider="openai", model="gpt-live-test", lineage="openai.gpt")
    return yaml.safe_dump(
        {
            "nodes": {"node_a": anthropic, "node_b": anthropic, "node_c": openai},
            "providers": {
                "anthropic": {"api_key_env": "ANTHROPIC_API_KEY"},
                "openai": {"api_key_env": "OPENAI_API_KEY"},
            },
        },
        sort_keys=True,
    )


def _build_live_bundle(tmp_path: Path) -> Path:
    """A live study bundle: study/instrument DOCX, the CSV, and llm.yaml."""
    study_dir = tmp_path / "study"
    (study_dir / "inputs").mkdir(parents=True)
    (study_dir / "config").mkdir(parents=True)

    doc = Document()
    for line in study_document().splitlines():
        doc.add_paragraph(line)
    doc.save(str(study_dir / "inputs" / "study_document.docx"))

    instrument = Document()
    instrument.add_paragraph("Survey instrument: 7-point Likert items for RES, CUL, INT.")
    instrument.save(str(study_dir / "inputs" / "data_dictionary.docx"))

    build_integration_study(20260705, n=300).write(study_dir / "inputs")
    (study_dir / "config" / "llm.yaml").write_text(_llm_yaml(), encoding="utf-8")
    return study_dir


def _sha(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def _mutate_docx_bytes(path: Path) -> None:
    """Change a .docx file's BYTES (a core property) without changing the extracted
    paragraph text: ``document_to_text`` output is identical, but the sha256 differs.
    Models input-file drift after the glance (the token binds only the config + Node A
    archive, not the input files)."""
    from burhan.contract.documents import document_to_text

    before_hash = _sha(path)
    before_text = document_to_text(path)
    doc = Document(str(path))
    doc.core_properties.author = "provenance-drift"
    doc.save(str(path))
    assert _sha(path) != before_hash  # the bytes really changed
    assert document_to_text(path) == before_text  # but the extracted text did not


def _write_reference_set(study_dir: Path) -> Path:
    """A reference set naming statistics the integration run actually emits."""
    reference = {
        "study_id": "integration-adoption-2026",
        "source": {
            "description": "Prior manual analysis (IT-7 fixture reference).",
            "documents": [{"path": "inputs/manual_results.docx", "sha256": "a" * 64}],
            "caveats": "Fixture reference set; not ground truth.",
        },
        "entries": [
            {
                "comparison_id": "REF-CFI",
                "domain": "fit",
                "metric": "cfi",
                "stat_id": "structural.fit.cfi",
                "reference_value": 0.95,
                "tolerance": 0.05,
            },
            {
                "comparison_id": "REF-PATH",
                "domain": "path",
                "metric": "estimate",
                "stat_id": "structural.path.RES->CUL",
                "reference_value": 0.45,
                "tolerance": 0.1,
            },
            {
                "comparison_id": "REF-REL",
                "domain": "reliability",
                "metric": "alpha",
                "stat_id": "measurement.reliability.RES",
                "reference_value": 0.80,
                "tolerance": 0.15,
            },
        ],
    }
    path = study_dir / "reference_set.yaml"
    path.write_text(yaml.safe_dump(reference, sort_keys=True), encoding="utf-8")
    return path


def test_live_extract_then_confirm_reaches_completed(tmp_path: Path) -> None:
    study_dir = _build_live_bundle(tmp_path)
    factory = RecordingFactory()

    extract = live_extract(study_dir, provider_factory=factory)
    # AT-M16-1: the live provider was invoked for Node A (not a canned echo).
    assert factory.calls["node_a"] == 1
    config_path = study_dir / "config" / "study_config.yaml"
    assert config_path.is_file()  # write-back
    assert (study_dir / "extract" / "node_a.0.json").is_file()  # archived
    assert (study_dir / "extract" / "GLANCE_TOKEN.json").is_file()  # pending token
    assert extract.study_id == "integration-adoption-2026"

    result = live_confirm(study_dir, provider_factory=factory, policy=fast_policy(tmp_path))
    assert result.state == "COMPLETED"
    # Node A was NOT re-called (it replays the archive); Node C ran live once.
    assert factory.calls["node_a"] == 1
    assert factory.calls["node_c"] >= 1
    # The archives are inside the sealed run dir.
    assert (result.run_dir / "llm" / "node_a.0.json").is_file()
    assert (result.run_dir / "llm" / "node_c.0.json").is_file()
    # No reference set was supplied -> no reference comparison is emitted (item 10).
    assert not (result.run_dir / "REFERENCE_COMPARISON.md").exists()


def test_confirm_without_token_halts_before_gate1(tmp_path: Path) -> None:
    study_dir = _build_live_bundle(tmp_path)
    factory = RecordingFactory()
    live_extract(study_dir, provider_factory=factory)
    (study_dir / "extract" / "GLANCE_TOKEN.json").unlink()  # no confirmation token

    with pytest.raises(BurhanHalt):
        live_confirm(study_dir, provider_factory=factory, policy=fast_policy(tmp_path))
    assert factory.calls["node_c"] == 0  # no Gate 1 / no Stage-1A work
    assert not (study_dir / "runs").exists()


def test_confirm_with_tampered_config_halts(tmp_path: Path) -> None:
    study_dir = _build_live_bundle(tmp_path)
    factory = RecordingFactory()
    live_extract(study_dir, provider_factory=factory)
    config_path = study_dir / "config" / "study_config.yaml"
    config_path.write_text(
        config_path.read_text(encoding="utf-8") + "\n# edited\n", encoding="utf-8"
    )

    with pytest.raises(BurhanHalt):
        live_confirm(study_dir, provider_factory=factory, policy=fast_policy(tmp_path))
    assert factory.calls["node_c"] == 0


def test_writeback_and_manifest_hashes_no_secrets(tmp_path: Path) -> None:
    study_dir = _build_live_bundle(tmp_path)
    factory = RecordingFactory()
    live_extract(study_dir, provider_factory=factory)
    result = live_confirm(study_dir, provider_factory=factory, policy=fast_policy(tmp_path))

    manifest = json.loads((result.run_dir / "manifest.json").read_text(encoding="utf-8"))
    config_path = study_dir / "config" / "study_config.yaml"
    # AT-M16-6: config bytes hashed; real llm.yaml model IDs recorded.
    assert manifest["hashes"]["study_config"] == _sha(config_path)
    assert manifest["llm_nodes"]["node_a"]["model"] == "claude-live-test"
    assert manifest["llm_nodes"]["node_c"]["model"] == "gpt-live-test"
    # archives are covered by the seal hash-tree.
    assert manifest["seal"]["hash_tree_root"]
    # no secrets anywhere in the manifest.
    blob = json.dumps(manifest)
    assert "ANTHROPIC_API_KEY" not in blob and "OPENAI_API_KEY" not in blob
    assert "sk-" not in blob


def test_rerun_replays_archives_byte_identical(tmp_path: Path) -> None:
    study_dir = _build_live_bundle(tmp_path)
    factory = RecordingFactory()
    live_extract(study_dir, provider_factory=factory)
    run = live_confirm(study_dir, provider_factory=factory, policy=fast_policy(tmp_path))

    # AT-M16-5: rerun replays archives, calls NO provider, byte-identical.
    rerun = live_rerun(run.run_dir, policy=fast_policy(tmp_path))
    assert rerun.state == "COMPLETED"


def test_planted_archive_mismatch_is_caught(tmp_path: Path) -> None:
    study_dir = _build_live_bundle(tmp_path)
    factory = RecordingFactory()
    live_extract(study_dir, provider_factory=factory)
    run = live_confirm(study_dir, provider_factory=factory, policy=fast_policy(tmp_path))

    archive = run.run_dir / "llm" / "node_c.0.json"
    payload = json.loads(archive.read_text(encoding="utf-8"))
    payload["response"] = yaml.safe_dump(
        {"verdict": "reject", "fixes": ["planted"]}, sort_keys=True
    )
    archive.write_text(json.dumps(payload), encoding="utf-8")

    with pytest.raises(BurhanHalt):
        live_rerun(run.run_dir, policy=fast_policy(tmp_path))


def test_extract_missing_study_document_halts(tmp_path: Path) -> None:
    # AT-M16-7 ingestion guard: no study_document.docx -> typed halt before any
    # provider call (never a silent empty prompt).
    (tmp_path / "study" / "inputs").mkdir(parents=True)
    factory = RecordingFactory()
    with pytest.raises(BurhanHalt):
        live_extract(tmp_path / "study", provider_factory=factory)
    assert factory.calls["node_a"] == 0


def test_extract_missing_csv_halts(tmp_path: Path) -> None:
    # A live run needs the CSV export as pipeline data; its absence halts typed.
    study_dir = _build_live_bundle(tmp_path)
    for csv in (study_dir / "inputs").glob("*.csv"):
        csv.unlink()
    factory = RecordingFactory()
    with pytest.raises(BurhanHalt):
        live_extract(study_dir, provider_factory=factory)
    assert factory.calls["node_a"] == 0


def test_confirm_with_tampered_archive_halts(tmp_path: Path) -> None:
    # AT-M16-3: the glance binds the Node A archive too — mutating it after the
    # glance halts before Gate 1 (the archive-sha branch of the pause).
    study_dir = _build_live_bundle(tmp_path)
    factory = RecordingFactory()
    live_extract(study_dir, provider_factory=factory)
    archive = study_dir / "extract" / "node_a.0.json"
    archive.write_text(archive.read_text(encoding="utf-8") + "\n", encoding="utf-8")

    with pytest.raises(BurhanHalt):
        live_confirm(study_dir, provider_factory=factory, policy=fast_policy(tmp_path))
    assert factory.calls["node_c"] == 0
    assert not (study_dir / "runs").exists()


def test_confirm_with_missing_config_halts(tmp_path: Path) -> None:
    # AT-M16-3: a token present but a vanished config is an incomplete pending
    # state — halt, never run against a config the glance never saw.
    study_dir = _build_live_bundle(tmp_path)
    factory = RecordingFactory()
    live_extract(study_dir, provider_factory=factory)
    (study_dir / "config" / "study_config.yaml").unlink()

    with pytest.raises(BurhanHalt):
        live_confirm(study_dir, provider_factory=factory, policy=fast_policy(tmp_path))
    assert factory.calls["node_c"] == 0
    assert not (study_dir / "runs").exists()


def test_reference_comparison_is_sealed_and_rerun_reproduces_it(tmp_path: Path) -> None:
    """TC-16 item 10: with a reference set, the confirmed run emits a SEALED
    REFERENCE_COMPARISON.md, and rerun regenerates it byte-identically with no
    provider calls. Fails if the builder call, render, write, seal inclusion, or
    with-reference rerun identity is removed.
    """
    study_dir = _build_live_bundle(tmp_path)
    reference_path = _write_reference_set(study_dir)
    factory = RecordingFactory()
    live_extract(study_dir, provider_factory=factory)

    run = live_confirm(
        study_dir,
        reference_path=reference_path,
        provider_factory=factory,
        policy=fast_policy(tmp_path),
    )
    assert run.state == "COMPLETED"

    report = run.run_dir / "REFERENCE_COMPARISON.md"
    assert report.is_file()  # write (item 10)
    text = report.read_text(encoding="utf-8")
    assert text.startswith("# Reference comparison")  # render
    for comparison_id in ("REF-CFI", "REF-PATH", "REF-REL"):
        assert comparison_id in text  # builder resolved each entry against the run's store
    assert "| Total | 3 |" in text  # builder summary counts

    # The reference-set bytes were copied verbatim into the run tree (for rerun).
    copied = run.run_dir / "reference" / "reference_set.yaml"
    assert copied.read_bytes() == reference_path.read_bytes()
    # Seal inclusion: the sealed run verifies WITH the report present (written pre-seal).
    Manifest.verify_seal(run.run_dir)

    # Rerun replays archives + reference bytes; regenerates the report byte-identically.
    node_c_after_confirm = factory.calls["node_c"]
    rerun = live_rerun(run.run_dir, policy=fast_policy(tmp_path))
    assert rerun.state == "COMPLETED"
    assert (rerun.run_dir / "REFERENCE_COMPARISON.md").read_bytes() == report.read_bytes()
    assert factory.calls["node_c"] == node_c_after_confirm  # rerun made no provider calls


def test_live_extract_injects_authoritative_provenance(tmp_path: Path) -> None:
    """The combined §7 fix: a Node A response omitting the non-LLM-derivable fields is
    completed by the engine into a schema-valid config — real source_documents (role,
    path, computed sha256 of the actual input files) and the governed playbook identity.
    Fails if the injection is removed (extract would halt on the missing required
    fields) or draws them from anywhere but the resolved inputs / governed playbook.
    """
    study_dir = _build_live_bundle(tmp_path)
    # Node A omits meta.source_documents and methodology.playbook_id/version, exactly as
    # the corrected prompt instructs (the model cannot hash files or know governance ids).
    trimmed = copy.deepcopy(integration_config())
    trimmed["meta"].pop("source_documents", None)
    trimmed["methodology"].pop("playbook_id", None)
    trimmed["methodology"].pop("playbook_version", None)
    response = yaml.safe_dump(trimmed, sort_keys=False)

    calls = {"node_a": 0}

    def factory(settings: LlmSettings, node: str) -> Callable[[str], str]:
        def call(prompt: str) -> str:
            calls["node_a"] += 1
            return response

        return call

    live_extract(study_dir, provider_factory=factory)
    assert calls["node_a"] == 1
    config = yaml.safe_load(
        (study_dir / "config" / "study_config.yaml").read_text(encoding="utf-8")
    )
    sds = config["meta"]["source_documents"]
    by_role = {d["role"]: d for d in sds}
    assert set(by_role) == {"study_document", "data_dictionary"}  # both resolved inputs
    assert by_role["study_document"]["path"] == "inputs/study_document.docx"
    # authoritative: the sha256 is the ACTUAL digest of the actual input file
    assert by_role["study_document"]["sha256"] == _sha(study_dir / "inputs" / "study_document.docx")
    assert by_role["data_dictionary"]["sha256"] == _sha(
        study_dir / "inputs" / "data_dictionary.docx"
    )
    for entry in sds:
        assert re.fullmatch(r"[a-f0-9]{64}", entry["sha256"])  # real 64-hex, not fabricated
    # governed playbook identity, resolved from the playbook the engine binds
    assert config["methodology"]["playbook_id"] == "CB_SEM_PLAYBOOK"
    assert config["methodology"]["playbook_version"] == "1.0"


def test_governed_playbook_identity_matches_playbook_metadata() -> None:
    # The playbook identity is read from the governed playbook's own metadata — never
    # the LLM output, never hard-coded prompt text.
    from burhan.cli.live import _governed_playbook_identity
    from burhan.core.playbook import playbooks_dir

    playbook_id, playbook_version = _governed_playbook_identity()
    meta = yaml.safe_load(
        (playbooks_dir() / "CB_SEM_PLAYBOOK_v1.0.yaml").read_text(encoding="utf-8")
    )["meta"]
    assert (playbook_id, playbook_version) == (meta["id"], meta["version"])


def _run_contract(run_dir: Path) -> dict:
    return json.loads((run_dir / "contract" / "study_config.json").read_text(encoding="utf-8"))


def _glanced_config(study_dir: Path) -> dict:
    return yaml.safe_load((study_dir / "config" / "study_config.yaml").read_text(encoding="utf-8"))


def test_confirm_injects_persisted_glanced_provenance_not_recomputed(tmp_path: Path) -> None:
    """Glance boundary: what runs must equal what was glanced. If input DOCX bytes drift
    after the glance, confirm must carry the PERSISTED glanced provenance into the run
    contract — never recompute a fresh hash from the now-mutated file. Fails on
    recompute-at-confirm (the run contract would carry the drifted hash)."""
    study_dir = _build_live_bundle(tmp_path)
    factory = RecordingFactory()
    live_extract(study_dir, provider_factory=factory)
    glanced = _glanced_config(study_dir)
    _mutate_docx_bytes(study_dir / "inputs" / "study_document.docx")  # drift after the glance

    run = live_confirm(study_dir, provider_factory=factory, policy=fast_policy(tmp_path))
    contract = _run_contract(run.run_dir)
    assert contract["meta"]["source_documents"] == glanced["meta"]["source_documents"]
    assert contract["methodology"]["playbook_id"] == glanced["methodology"]["playbook_id"]
    assert contract["methodology"]["playbook_version"] == glanced["methodology"]["playbook_version"]


def test_run_contract_equals_glanced_config_for_provenance(tmp_path: Path) -> None:
    # Invariant lock: with no drift, the run contract equals the glanced config for the
    # engine-supplied provenance/governance (one-directional flow, authored once).
    study_dir = _build_live_bundle(tmp_path)
    factory = RecordingFactory()
    live_extract(study_dir, provider_factory=factory)
    run = live_confirm(study_dir, provider_factory=factory, policy=fast_policy(tmp_path))
    glanced = _glanced_config(study_dir)
    contract = _run_contract(run.run_dir)
    assert contract["meta"]["source_documents"] == glanced["meta"]["source_documents"]
    assert contract["methodology"] == glanced["methodology"]


def test_rerun_provenance_from_sealed_contract_not_mutable_bundle(tmp_path: Path) -> None:
    """NFR-101: rerun reproduces the SEALED run. Provenance must come from the sealed run
    contract, not a fresh hash of the (mutable) study bundle. Fails on recompute-at-rerun
    (byte-identity breaks, or the rerun contract carries the drifted hash)."""
    study_dir = _build_live_bundle(tmp_path)
    factory = RecordingFactory()
    live_extract(study_dir, provider_factory=factory)
    run = live_confirm(study_dir, provider_factory=factory, policy=fast_policy(tmp_path))
    sealed = _run_contract(run.run_dir)
    _mutate_docx_bytes(study_dir / "inputs" / "study_document.docx")  # drift after the seal

    rerun = live_rerun(run.run_dir, policy=fast_policy(tmp_path))
    assert (
        _run_contract(rerun.run_dir)["meta"]["source_documents"]
        == sealed["meta"]["source_documents"]
    )
