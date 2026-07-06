"""AT-M16-7: DOCX->text ingestion for Node A's text inputs (TC-16).

``document_to_text(path)`` yields the study/instrument text Node A consumes;
a corrupt, missing, or empty DOCX halts typed (``IntegrityHalt``) — never
silently empty. The raw CSV is never a document and never passes through here
(that boundary is proven separately in ``test_live_boundary``).
"""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document  # type: ignore[import-untyped]

from burhan.contract.documents import document_to_text
from burhan.core.errors import IntegrityHalt


def _make_docx(path: Path, paragraphs: list[str], table: list[list[str]] | None = None) -> None:
    doc = Document()
    for para in paragraphs:
        doc.add_paragraph(para)
    if table:
        built = doc.add_table(rows=len(table), cols=len(table[0]))
        for r, row in enumerate(table):
            for c, cell in enumerate(row):
                built.rows[r].cells[c].text = cell
    doc.save(str(path))


def test_docx_to_text_extracts_paragraphs_and_tables_in_order(tmp_path: Path) -> None:
    docx_path = tmp_path / "study.docx"
    _make_docx(
        docx_path,
        paragraphs=["Study of organizational readiness", "Second paragraph"],
        table=[["Item", "Construct"], ["Q1", "TI"]],
    )
    text = document_to_text(docx_path)
    assert "Study of organizational readiness" in text
    assert "Second paragraph" in text
    assert "Q1" in text and "TI" in text
    # document order: the opening paragraph precedes the table cell
    assert text.index("Study of organizational readiness") < text.index("Q1")


def test_corrupt_docx_halts_typed(tmp_path: Path) -> None:
    bad = tmp_path / "broken.docx"
    bad.write_text("this is not a docx package", encoding="utf-8")
    with pytest.raises(IntegrityHalt):
        document_to_text(bad)


def test_missing_docx_halts_typed(tmp_path: Path) -> None:
    with pytest.raises(IntegrityHalt):
        document_to_text(tmp_path / "does_not_exist.docx")


def test_empty_docx_halts_never_silently_empty(tmp_path: Path) -> None:
    empty = tmp_path / "empty.docx"
    _make_docx(empty, paragraphs=[])
    with pytest.raises(IntegrityHalt):
        document_to_text(empty)
